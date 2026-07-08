from __future__ import annotations

import json
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
PROJECT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "output"
FIGURES = PROJECT / "paper" / "figures"
TEMPLATE = ROOT / "20260428-V3.docx"
MANIFEST = PROJECT / "data" / "processed" / "manifest.json"
DOCX_OUT = OUTPUT / "论文重大修改稿-EcoSpec-KG-V1.docx"

TITLE = "面向生态评估技术规范数字化的模式约束与证据溯源知识图谱构建及补全方法"
EN_TITLE = (
    "Schema-Constrained and Evidence-Traceable Knowledge Graph Construction "
    "and Completion for Digital Ecological Assessment Standards"
)
PLACEHOLDER = "待实验补充"

FONT_CN = "宋体"
FONT_HEADING = "黑体"
FONT_EN = "Times New Roman"
FONT_PATH = Path(r"C:\Windows\Fonts\msyh.ttc")


def set_run_font(run, cn=FONT_CN, en=FONT_EN, size=10.5, bold=False, color=None):
    run.font.name = en
    run._element.rPr.rFonts.set(qn("w:eastAsia"), cn)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def clear_body(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = FONT_EN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.35
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.first_line_indent = Pt(21)

    for name, size, before, after in (
        ("Heading 1", 14, 10, 5),
        ("Heading 2", 12, 8, 4),
        ("Heading 3", 10.5, 6, 3),
    ):
        style = doc.styles[name]
        style.font.name = FONT_EN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_HEADING)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.first_line_indent = Pt(0)
        style.paragraph_format.keep_with_next = True

    caption = doc.styles["Caption"]
    caption.font.name = FONT_EN
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    caption.font.size = Pt(9)
    caption.font.color.rgb = RGBColor(0, 0, 0)
    caption.paragraph_format.first_line_indent = Pt(0)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if "参考文献" in doc.styles:
        ref = doc.styles["参考文献"]
        ref.font.name = FONT_EN
        ref._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        ref.font.size = Pt(9)
        ref.paragraph_format.first_line_indent = Pt(-18)
        ref.paragraph_format.left_indent = Pt(18)
        ref.paragraph_format.line_spacing = 1.0


def add_paragraph(
    doc: Document,
    text: str,
    style: str | None = None,
    *,
    align=None,
    bold=False,
    size=None,
    cn_font=FONT_CN,
    first_indent=True,
):
    paragraph = doc.add_paragraph(style=style)
    if align is not None:
        paragraph.alignment = align
    if not first_indent:
        paragraph.paragraph_format.first_line_indent = Pt(0)
    run = paragraph.add_run(text)
    set_run_font(
        run,
        cn=cn_font,
        size=size or (10.5 if style is None else run.font.size.pt if run.font.size else 10.5),
        bold=bold,
    )
    return paragraph


def add_mixed_paragraph(doc: Document, parts: list[tuple[str, bool]], style=None):
    paragraph = doc.add_paragraph(style=style)
    for text, bold in parts:
        run = paragraph.add_run(text)
        set_run_font(run, bold=bold)
    return paragraph


def set_cell_width(cell, width_cm: float) -> None:
    width = Cm(width_cm)
    cell.width = width
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width.twips)))
    tc_w.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), "808080")


def set_cell_text(cell, text: str, bold=False, size=8.5, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(text))
    set_run_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge in ("top", "left", "bottom", "right"):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), "90")
        node.set(qn("w:type"), "dxa")


def add_table(doc: Document, headers, rows, widths, caption: str):
    caption_paragraph = add_paragraph(
        doc,
        caption,
        style="Caption",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
        size=9,
        first_indent=False,
    )
    caption_paragraph.paragraph_format.keep_with_next = True
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    for index, (header, width) in enumerate(zip(headers, widths)):
        cell = table.rows[0].cells[index]
        set_cell_width(cell, width)
        set_cell_text(cell, header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(cell, "D9E2F3")
    for row in rows:
        cells = table.add_row().cells
        for index, (value, width) in enumerate(zip(row, widths)):
            set_cell_width(cells[index], width)
            align = (
                WD_ALIGN_PARAGRAPH.CENTER
                if len(str(value)) <= 16
                else WD_ALIGN_PARAGRAPH.LEFT
            )
            set_cell_text(cells[index], value, align=align)
            if str(value) == PLACEHOLDER:
                shade_cell(cells[index], "FFF2CC")
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def draw_arrow(draw, start, end, color=(70, 82, 92), width=4):
    draw.line([start, end], fill=color, width=width)
    x2, y2 = end
    draw.polygon(
        [(x2, y2), (x2 - 14, y2 - 8), (x2 - 14, y2 + 8)], fill=color
    )


def draw_flow(path: Path, title: str, rows: list[list[tuple[str, str]]]):
    width, height = 1800, 300 + len(rows) * 250
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = ImageFont.truetype(str(FONT_PATH), 48)
    node_font = ImageFont.truetype(str(FONT_PATH), 31)
    small_font = ImageFont.truetype(str(FONT_PATH), 24)
    draw.text((70, 45), title, font=title_font, fill=(25, 35, 45))
    palette = {
        "blue": (225, 237, 247),
        "green": (225, 241, 232),
        "yellow": (252, 241, 207),
        "gray": (235, 238, 240),
        "red": (248, 228, 225),
    }
    for row_index, row in enumerate(rows):
        y = 150 + row_index * 250
        box_w = 280
        gap = (width - 140 - len(row) * box_w) // max(1, len(row) - 1)
        for index, (label, color_key) in enumerate(row):
            x = 70 + index * (box_w + gap)
            draw.rounded_rectangle(
                [x, y, x + box_w, y + 120],
                radius=8,
                fill=palette[color_key],
                outline=(83, 96, 106),
                width=2,
            )
            bbox = draw.multiline_textbbox((0, 0), label, font=node_font, spacing=6)
            tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
            draw.multiline_text(
                (x + (box_w - tw) / 2, y + (120 - th) / 2),
                label,
                font=node_font,
                fill=(20, 30, 40),
                align="center",
                spacing=6,
            )
            if index < len(row) - 1:
                draw_arrow(
                    draw,
                    (x + box_w + 10, y + 60),
                    (x + box_w + gap - 10, y + 60),
                )
        if row_index < len(rows) - 1:
            draw.text(
                (75, y + 150),
                "↓  数据与证据沿统一Schema传递",
                font=small_font,
                fill=(70, 82, 92),
            )
    path.parent.mkdir(parents=True, exist_ok=True)
    image.save(path, dpi=(180, 180))


def create_figures() -> list[Path]:
    figures = [
        FIGURES / "fig1_research_framework.png",
        FIGURES / "fig2_dataset_pipeline.png",
        FIGURES / "fig3_method_architecture.png",
        FIGURES / "fig4_case_path.png",
    ]
    draw_flow(
        figures[0],
        "EcoSpec-KG研究框架",
        [
            [
                ("规范语料\nHJ 1166-1176", "gray"),
                ("结构化解析\n正文/表格/公式", "blue"),
                ("专家标注\n双人独立复核", "green"),
                ("分组划分\n防止路径泄漏", "yellow"),
            ],
            [
                ("双GraphRAG\n索引与检索", "blue"),
                ("Schema约束\n关系合法性", "green"),
                ("LoRA补全\n低资源适配", "yellow"),
                ("证据拒绝\n条款可追溯", "red"),
            ],
            [
                ("对比实验", "gray"),
                ("消融实验", "gray"),
                ("专家验证", "gray"),
                ("规范检索原型", "blue"),
            ],
        ],
    )
    draw_flow(
        figures[1],
        "生态评估规范数据集构建流程",
        [
            [
                ("文件哈希去重", "gray"),
                ("标准编号核验", "blue"),
                ("分页与章节识别", "blue"),
                ("可追溯文本块", "green"),
            ],
            [
                ("模型候选抽取", "yellow"),
                ("标注指南约束", "green"),
                ("两名专家复核", "green"),
                ("共识金标准", "blue"),
            ],
            [
                ("训练集70%", "gray"),
                ("验证集15%", "gray"),
                ("内部测试15%", "gray"),
                ("跨规范测试集", "yellow"),
            ],
        ],
    )
    draw_flow(
        figures[2],
        "模式约束与证据溯源方法架构",
        [
            [
                ("原生适配器\n轻量可复现", "blue"),
                ("统一图模型", "green"),
                ("官方适配器\nGraphRAG 3.1", "blue"),
            ],
            [
                ("候选实体关系", "yellow"),
                ("Schema类型校验", "green"),
                ("LoRA关系补全", "yellow"),
                ("证据原文匹配", "red"),
            ],
            [
                ("接受并入图", "green"),
                ("保留审核状态", "gray"),
                ("拒绝无证据关系", "red"),
            ],
        ],
    )
    draw_flow(
        figures[3],
        "水源涵养指标的可追溯知识路径示例",
        [
            [
                ("水源涵养量\n评估指标", "blue"),
                ("水量平衡方程\n计算公式", "green"),
                ("降雨/径流/蒸散\n模型参数", "yellow"),
            ],
            [
                ("遥感与地面数据\n数据来源", "gray"),
                ("单位与时空范围", "gray"),
                ("质量控制要求", "red"),
                ("HJ 1173条款\n原文证据", "blue"),
            ],
        ],
    )
    return figures


def add_figure(doc: Document, path: Path, caption_cn: str, caption_en: str):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(6.25))
    add_paragraph(
        doc,
        caption_cn,
        style="Caption",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=9,
        first_indent=False,
    )
    add_paragraph(
        doc,
        caption_en,
        style="Caption",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=8.5,
        first_indent=False,
    )


def add_page_number(section) -> None:
    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    set_run_font(run, size=9)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


def build_document() -> Path:
    OUTPUT.mkdir(parents=True, exist_ok=True)
    figures = create_figures()
    doc = Document(TEMPLATE)
    clear_body(doc)
    configure_styles(doc)
    section = doc.sections[0]
    section.top_margin = Cm(2.3)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)
    add_page_number(section)

    p = add_paragraph(
        doc,
        TITLE,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
        size=18,
        cn_font=FONT_HEADING,
        first_indent=False,
    )
    p.paragraph_format.space_after = Pt(8)
    add_paragraph(
        doc,
        "陈艳¹，沈志龙²，张京¹，蒋娴¹*",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=11,
        first_indent=False,
    )
    add_paragraph(
        doc,
        "（1. 中国林业科学研究院资源信息研究所，北京 100091；"
        "2. 杭州鑫宥算法科技有限公司，杭州 310051）",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=9.5,
        first_indent=False,
    )

    add_mixed_paragraph(
        doc,
        [
            ("摘要：", True),
            (
                "全国生态状况调查评估技术规范已形成较完整的指标、模型、数据和质量控制体系，"
                "但其知识分布于正文、表格、公式及附录，对人工业务具有可读性，却缺少统一的机器可读表达。"
                "本文面向规范数字化而非生态模型替代，提出EcoSpec-KG框架。该框架以HJ 1166-HJ 1176"
                "系列标准为语料，通过文件去重、分页分节和知识路径分组构建可追溯数据集；设计覆盖指标、"
                "公式、参数、数据源、时空范围和质量要求的领域Schema；在统一图模型上同时实现轻量原生"
                "GraphRAG与Microsoft GraphRAG适配器，并利用Qwen3-0.6B和LoRA完成低资源关系补全。"
                "所有候选关系均需通过类型约束和原文证据匹配，无证据关系被拒绝入图。实验设置包括规则抽取、"
                "大模型提示、向量RAG、两类GraphRAG、Schema约束和完整方法的对比，并设计组件消融、"
                "跨规范测试和双专家盲评。当前稿为实验实施前版本，性能结果及专家统计均待正式标注与GPU实验后补充，"
                "不据此形成性能结论。研究预期为生态评估规范知识检索、参数依赖检查和标准更新影响分析提供"
                "可复用的数据与技术基础。",
                False,
            ),
        ],
    )
    add_mixed_paragraph(
        doc,
        [
            ("关键词：", True),
            ("生态评估技术规范；知识图谱；GraphRAG；证据溯源；知识补全；标准数字化", False),
        ],
    )
    add_paragraph(
        doc,
        EN_TITLE,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
        size=14,
        first_indent=False,
    )
    add_paragraph(
        doc,
        "CHEN Yan¹, SHEN Zhilong², ZHANG Jing¹, JIANG Xian¹*",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=10,
        first_indent=False,
    )
    add_paragraph(
        doc,
        "(1. Institute of Resource Information, Chinese Academy of Forestry, "
        "Beijing 100091, China; 2. Hangzhou Xinyou Algorithm Technology Co., "
        "Ltd., Hangzhou 310051, China)",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=9,
        first_indent=False,
    )
    add_mixed_paragraph(
        doc,
        [
            ("Abstract: ", True),
            (
                "China's technical specifications for national ecological status investigation "
                "and assessment provide established indicators, models, data requirements, and "
                "quality-control procedures. However, knowledge is distributed across narrative "
                "clauses, tables, equations, and appendices, which limits unified machine-readable "
                "access. This study proposes EcoSpec-KG for standards digitization rather than "
                "replacement of geospatial or ecosystem models. Eleven standards from HJ 1166 to "
                "HJ 1176 are deduplicated and segmented with page- and clause-level provenance. "
                "A domain schema constrains indicators, formulas, parameters, data sources, "
                "spatiotemporal scope, and quality requirements. A shared graph model supports "
                "both a reproducible native GraphRAG pipeline and a Microsoft GraphRAG adapter, "
                "while Qwen3-0.6B with LoRA is reserved for low-resource relation completion. "
                "Every accepted relation must satisfy type constraints and match source evidence. "
                "Comparative, ablation, cross-standard, and two-expert evaluations are specified. "
                "As this is a pre-experiment manuscript, quantitative results remain to be filled "
                "after annotation and GPU runs; no performance claim is made at this stage.",
                False,
            ),
        ],
    )
    add_mixed_paragraph(
        doc,
        [
            ("Keywords: ", True),
            (
                "ecological assessment standard; knowledge graph; GraphRAG; provenance; "
                "knowledge graph completion; standards digitization",
                False,
            ),
        ],
    )

    add_paragraph(doc, "1 引言", style="Heading 1", first_indent=False)
    add_paragraph(
        doc,
        "生态系统质量与服务功能评估通常依赖遥感、气象、水文、土壤和野外调查数据，"
        "其核心计算仍由地理空间分析和生态模型承担。HJ 1172-2021和HJ 1173-2021分别规定了"
        "生态系统质量与服务功能评估的指标体系、技术流程和主要计算方法[1-2]；HJ 1166-2021"
        "及HJ 1176-2021进一步规定了遥感解译、野外核查和数据质量控制要求[3-4]。因此，"
        "本文不把现行规范描述为业务逻辑不完整，也不把知识图谱作为生态模型的替代方案。",
    )
    add_paragraph(
        doc,
        "实际信息化建设面临的是另一类问题：规范面向专业人员编写，指标定义、公式、参数说明、"
        "数据来源、适用条件和质量要求分散在多项标准及其附录中。技术人员能够通过人工阅读完成梳理，"
        "但软件系统难以在缺少统一语义类型和条款溯源的情况下自动建立依赖路径。该问题属于规范知识"
        "组织与机器可读化，而非生态评估理论或计算模型本身的缺陷。",
    )
    add_paragraph(
        doc,
        "知识图谱适合表达实体、关系和约束，已被用于复杂领域知识组织[5-7]。RAG通过外部知识检索"
        "增强生成过程[8]，GraphRAG进一步利用实体图和社区结构组织跨段落信息[9]。但通用GraphRAG"
        "强调面向问答的实体发现与社区摘要，不能直接保证规范中的关系类型合法、公式方向正确或输出"
        "具有可核验条款证据。与此同时，低资源知识补全仍面临标注规模有限和模型幻觉问题。",
    )
    add_paragraph(
        doc,
        "针对上述问题，本文提出EcoSpec-KG。其贡献包括：第一，构建覆盖HJ 1166-HJ 1176的"
        "可追溯规范语料及标注方案；第二，设计生态评估规范专用Schema和完整知识路径分组策略；"
        "第三，在统一图模型上提供原生与官方双GraphRAG适配，并引入Schema校验和原文证据拒绝机制；"
        "第四，建立对比、消融、跨规范和双专家验证体系，评价知识抽取与补全质量，而不是以图密度替代"
        "业务正确性。整体研究框架见图1。",
    )
    add_figure(
        doc,
        figures[0],
        "图1 EcoSpec-KG研究框架",
        "Fig. 1 Research framework of EcoSpec-KG",
    )

    add_paragraph(doc, "2 规范语料与问题定义", style="Heading 1", first_indent=False)
    add_paragraph(doc, "2.1 语料范围与版本控制", style="Heading 2", first_indent=False)
    add_paragraph(
        doc,
        "原始语料来自全国生态状况调查评估技术规范文件夹。系统使用SHA-256识别完全重复文件，"
        "再按标准编号聚合内容不同但主题相同的副本，每项标准仅保留一个正式PDF。DOCX副本不参与训练，"
        "2020年编制说明仅用于研究背景，不作为规范性知识。经去重，语料包含HJ 1166-HJ 1176共11项"
        "正式标准，累计分块129个。正式数据集规模将在完成专家标注后由程序自动统计。",
    )
    manifest = json.loads(MANIFEST.read_text(encoding="utf-8"))
    role_map = {"core": "核心语料", "support": "支撑语料", "external_test": "跨规范测试"}
    table1 = [
        (
            item["standard_code"],
            item["title"],
            role_map[item["role"]],
            item["page_count"],
        )
        for item in manifest
    ]
    add_table(
        doc,
        ["标准编号", "规范主题", "数据角色", "页数"],
        table1,
        [3.1, 7.1, 3.2, 1.4],
        "表1 去重后的规范语料构成",
    )
    add_paragraph(doc, "2.2 研究任务", style="Heading 2", first_indent=False)
    add_paragraph(
        doc,
        "本文设置三个相互关联的任务。任务一是规范实体与关系抽取，即从带页码和章节标识的文本块中"
        "识别指标、公式、参数和数据要求。任务二是知识补全，即在候选实体范围内预测缺失关系，同时"
        "要求补全结果满足Schema。任务三是证据溯源，即为每条关系返回标准编号、页码、章节和原文片段，"
        "并拒绝无法在来源文本中匹配的关系。",
    )
    add_paragraph(doc, "2.3 研究边界", style="Heading 2", first_indent=False)
    add_paragraph(
        doc,
        "本文数据仅包含规范文本，不包含真实区域的遥感影像、气象、土壤或野外样方数据。因而实验"
        "评价的是标准知识的抽取、补全、检索与溯源能力，不评价水源涵养量、土壤保持量或EQI的实际"
        "计算精度。知识图谱输出可以为后续模型配置提供参数清单，但不能直接视为区域生态评估结果。",
    )

    add_paragraph(doc, "3 数据集与领域Schema", style="Heading 1", first_indent=False)
    add_paragraph(doc, "3.1 文档解析与分块", style="Heading 2", first_indent=False)
    add_paragraph(
        doc,
        "系统按页提取文本并保留文件哈希、页码和章节。默认文本块长度为1400字符、重叠150字符；"
        "正式实验同时测试150、300和600 token三种分块尺度。每个DocumentChunk包含标准编号、主题、"
        "页码、章节、块类型、正文、来源文件和知识路径标识。路径标识用于保证同一指标链不会同时"
        "进入训练集和测试集。",
    )
    add_figure(
        doc,
        figures[1],
        "图2 生态评估规范数据集构建流程",
        "Fig. 2 Dataset construction workflow for ecological assessment standards",
    )
    add_paragraph(doc, "3.2 实体和关系模式", style="Heading 2", first_indent=False)
    entity_rows = [
        ("评估指标", "indicator", "水源涵养量、生态系统质量"),
        ("计算公式", "formula", "水量平衡方程、RUSLE、EQI"),
        ("模型参数", "parameter", "降雨量、地表径流量、FVC"),
        ("单位", "unit", "mm/a、t/(hm²·a)"),
        ("数据来源", "data_source", "气象站数据、遥感影像"),
        ("处理方法", "method", "样方调查、插值、归一化"),
        ("生态系统类型", "ecosystem_type", "森林、草地、荒漠"),
        ("空间范围", "spatial_scope", "全国、省级行政区、评估区"),
        ("时间范围", "temporal_scope", "一年、多年、基准年"),
        ("质量控制要求", "quality_requirement", "精度、时相、完整性要求"),
        ("标准条款", "standard_clause", "HJ编号、章节和附录"),
    ]
    add_table(
        doc,
        ["实体类型", "代码", "示例"],
        entity_rows,
        [3.2, 4.0, 7.6],
        "表2 生态评估规范实体Schema",
    )
    relation_rows = [
        ("采用公式", "uses_formula", "评估指标→计算公式"),
        ("依赖参数", "depends_on", "公式/指标→参数"),
        ("数据来源于", "sourced_from", "参数→数据来源"),
        ("单位为", "has_unit", "参数/指标→单位"),
        ("适用于", "applies_to", "指标/公式→生态类型或时空范围"),
        ("定义于条款", "defined_in", "知识实体→标准条款"),
        ("采用处理方法", "uses_method", "指标/参数→处理方法"),
        ("受质量要求约束", "constrained_by", "数据或参数→质量要求"),
        ("由参数派生", "derived_from", "指标/参数→基础参数"),
    ]
    add_table(
        doc,
        ["关系类型", "代码", "允许方向"],
        relation_rows,
        [3.5, 4.4, 6.9],
        "表3 关系类型与方向约束",
    )
    add_paragraph(doc, "3.3 人工标注与数据划分", style="Heading 2", first_indent=False)
    add_paragraph(
        doc,
        "候选关系由规则基线和大模型生成，但自动结果不直接作为金标准。两名专家依据统一指南独立"
        "判断实体边界、关系方向、生态逻辑和证据充分性；分歧通过共识讨论解决。数据按完整path_id"
        "分组后以70%、15%和15%划分训练集、验证集和内部测试集；HJ 1171、HJ 1174和HJ 1175"
        "整体作为跨规范测试集。最终实体数、关系数和一致性统计均为待实验补充。",
    )

    add_paragraph(doc, "4 EcoSpec-KG方法", style="Heading 1", first_indent=False)
    add_paragraph(doc, "4.1 双GraphRAG适配", style="Heading 2", first_indent=False)
    add_paragraph(
        doc,
        "为区分通用框架能力与本文约束机制，系统实现两个可交换适配器。原生适配器使用统一文本块、"
        "倒排与向量检索、实体关系图和社区结构，强调轻量、可复现和便于消融；官方适配器固定使用"
        "Microsoft GraphRAG 3.1.0，通过OpenAI兼容接口连接远程Qwen模型，并将其Parquet输出映射"
        "到同一数据类型。两类适配器共享语料、Schema、数据划分和评测接口。",
    )
    add_paragraph(doc, "4.2 Schema约束抽取", style="Heading 2", first_indent=False)
    add_paragraph(
        doc,
        "模型首先从候选文本块输出结构化实体关系。每个预测必须给出头尾实体类型和关系类型，系统"
        "依据表3执行合法性检查。例如，uses_formula仅允许“评估指标→计算公式”，单位实体不能作为"
        "公式主体。该机制把部分业务规则从生成提示转移到确定性校验层，避免仅依赖自然语言提示。",
    )
    add_paragraph(doc, "4.3 低资源关系补全", style="Heading 2", first_indent=False)
    add_paragraph(
        doc,
        "补全模型采用Qwen3-0.6B[10]，结构化抽取关闭思考模式。LoRA[11]秩设为8，alpha为16，"
        "dropout为0.05，学习率为2×10⁻⁴，最多训练5轮并使用早停。候选实体由语义相似度Top-10"
        "生成，默认阈值0.75。正式实验分别使用42、43和44三个随机种子，并完整保存运行清单。",
    )
    add_paragraph(doc, "4.4 原文证据拒绝机制", style="Heading 2", first_indent=False)
    add_paragraph(
        doc,
        "关系通过Schema后仍不能直接入图。系统对证据文本执行空白归一化，并要求其为指定"
        "DocumentChunk的精确子串，同时核对标准编号和文本块标识。无法定位证据的关系被标记为"
        "unsupported并拒绝入图。该机制不保证关系必然正确，但能够区分“有规范证据的候选关系”"
        "与“模型自由生成内容”，为专家复核提供可审计入口。",
    )
    add_figure(
        doc,
        figures[2],
        "图3 模式约束与证据溯源方法架构",
        "Fig. 3 Architecture of schema constraints and evidence tracing",
    )

    add_paragraph(doc, "5 实验设计", style="Heading 1", first_indent=False)
    add_paragraph(doc, "5.1 对比方法", style="Heading 2", first_indent=False)
    method_rows = [
        ("Rule", "规则模板抽取", "不使用大模型"),
        ("Zero-shot", "Qwen直接抽取", "无示例"),
        ("Few-shot", "Qwen上下文学习", "3个示例"),
        ("Vector RAG", "向量检索后抽取", "无图结构"),
        ("MS GraphRAG", "官方GraphRAG 3.1.0", "统一模型接口"),
        ("Native GraphRAG", "原生图检索适配器", "无Schema"),
        ("Schema GraphRAG", "加入类型约束", "无LoRA"),
        ("EcoSpec-KG", "Schema+GraphRAG+LoRA+证据校验", "完整方法"),
    ]
    add_table(
        doc,
        ["方法", "主要设置", "用途"],
        method_rows,
        [3.4, 7.0, 4.4],
        "表4 对比方法设置",
    )
    add_paragraph(doc, "5.2 评价指标", style="Heading 2", first_indent=False)
    add_paragraph(
        doc,
        "实体识别和关系抽取报告Precision、Recall、Micro-F1和Macro-F1；链接预测报告MRR、"
        "Hits@1、Hits@3和Hits@10；业务完整性报告“指标—公式—参数—数据源—质量要求”路径"
        "准确率；证据质量报告条款溯源准确率和无证据率。弱连通分量、最大连通分量和图密度仅用于"
        "描述图结构变化，不作为生态知识正确性的替代指标。",
    )
    add_paragraph(doc, "5.3 主实验与消融实验", style="Heading 2", first_indent=False)
    result_rows = [
        (name, PLACEHOLDER, PLACEHOLDER, PLACEHOLDER, PLACEHOLDER, PLACEHOLDER)
        for name in (
            "Rule",
            "Zero-shot",
            "Few-shot",
            "Vector RAG",
            "MS GraphRAG",
            "Native GraphRAG",
            "Schema GraphRAG",
            "EcoSpec-KG",
        )
    ]
    add_table(
        doc,
        ["方法", "Micro-F1", "Macro-F1", "MRR", "路径准确率", "溯源准确率"],
        result_rows,
        [3.3, 2.5, 2.5, 2.2, 2.9, 2.9],
        "表5 主实验结果（正式实验后补充）",
    )
    ablation_rows = [
        ("完整方法", PLACEHOLDER, PLACEHOLDER, PLACEHOLDER),
        ("去除Schema", PLACEHOLDER, PLACEHOLDER, PLACEHOLDER),
        ("去除GraphRAG", PLACEHOLDER, PLACEHOLDER, PLACEHOLDER),
        ("去除LoRA", PLACEHOLDER, PLACEHOLDER, PLACEHOLDER),
        ("去除few-shot", PLACEHOLDER, PLACEHOLDER, PLACEHOLDER),
        ("去除证据校验", PLACEHOLDER, PLACEHOLDER, PLACEHOLDER),
    ]
    add_table(
        doc,
        ["设置", "关系F1", "路径准确率", "无证据率"],
        ablation_rows,
        [5.2, 3.3, 3.6, 3.6],
        "表6 消融实验结果（正式实验后补充）",
    )
    add_paragraph(
        doc,
        "除组件消融外，参数敏感性实验比较150、300和600 token分块，0/1/3/5-shot提示，"
        "以及0.65、0.75和0.85候选阈值。所有主实验以三个随机种子重复运行，报告均值与标准差。"
        "当前代码仅完成CPU烟雾测试，表5和表6不得据此填写。",
    )

    add_paragraph(doc, "6 专家验证与规范检索案例", style="Heading 1", first_indent=False)
    add_paragraph(doc, "6.1 双专家盲评", style="Heading 2", first_indent=False)
    add_paragraph(
        doc,
        "从模型新增关系中按关系类型和标准分层抽取120条样本，由两名具有生态评估或遥感工作经验的"
        "专家独立评价事实正确性、方向正确性、生态逻辑和证据充分性。每项采用三级评分，并额外记录"
        "是否建议入图。报告原始一致性和共识后接受率；若最终仅有一名专家，则只报告单人复核结果，"
        "不计算一致性统计。",
    )
    expert_rows = [
        ("事实正确性", "三级评分", PLACEHOLDER),
        ("关系方向", "三级评分", PLACEHOLDER),
        ("生态逻辑", "三级评分", PLACEHOLDER),
        ("证据充分性", "三级评分", PLACEHOLDER),
        ("入图接受率", "接受/拒绝", PLACEHOLDER),
        ("专家一致性", "Cohen's κ", PLACEHOLDER),
    ]
    add_table(
        doc,
        ["评价维度", "记录方式", "结果"],
        expert_rows,
        [5.0, 5.0, 5.7],
        "表7 专家验证结果（专家复核后补充）",
    )
    add_paragraph(doc, "6.2 规范检索案例", style="Heading 2", first_indent=False)
    add_paragraph(
        doc,
        "案例覆盖水源涵养、土壤保持、防风固沙和生态系统质量指数。系统输出指标采用的公式、公式"
        "依赖参数、参数数据来源、单位、时空范围、质量要求及对应条款。正式验证将设置24项检索任务，"
        "比较人工PDF检索、向量RAG和EcoSpec-KG的答案正确率、证据准确率和完成时间。结果待实验补充。",
    )
    add_figure(
        doc,
        figures[3],
        "图4 水源涵养指标的可追溯知识路径示例",
        "Fig. 4 Traceable knowledge path example for water conservation",
    )
    add_paragraph(doc, "6.3 机器可执行性原型", style="Heading 2", first_indent=False)
    add_paragraph(
        doc,
        "原型以Gradio提供规范检索、指标依赖路径、缺失项检查和原文证据四个工作区。它可以生成"
        "公式依赖树和输入检查清单，但不加载真实区域生态数据，也不输出生态质量等级。该原型用于"
        "验证知识组织结果能否转化为后续软件接口，而非证明生态计算精度。",
    )

    add_paragraph(doc, "7 讨论", style="Heading 1", first_indent=False)
    add_paragraph(doc, "7.1 预期意义", style="Heading 2", first_indent=False)
    add_paragraph(
        doc,
        "本文价值主要体现在规范数字化基础设施。条款级溯源可以降低知识抽取结果的审核成本，Schema"
        "可以在生成模型之外提供确定性业务约束，完整路径表示可以支持指标参数清单生成、数据准备检查"
        "和标准修订影响分析。公开数据结构和双适配器接口有利于后续研究在相同语料和划分上复现比较。",
    )
    add_paragraph(doc, "7.2 与实际生态评估的关系", style="Heading 2", first_indent=False)
    add_paragraph(
        doc,
        "知识图谱回答“规范要求什么”和“知识位于何处”，而遥感、GIS和过程模型回答“特定区域的"
        "生态状态如何”。前者可以为后者提供配置和审查依据，但不能替代数据预处理、空间尺度统一、"
        "模型参数率定和不确定性分析。将机器可读路径表述为真实生态评估结果会超出本文证据范围。",
    )
    add_paragraph(doc, "7.3 局限性", style="Heading 2", first_indent=False)
    add_paragraph(
        doc,
        "第一，语料限定为2021年发布的系列标准，后续修改单和地方规范需要增量处理。第二，PDF公式"
        "和复杂表格的文本抽取仍可能发生符号丢失，需结合人工复核或版面模型。第三，两名专家能够提供"
        "基本一致性信息，但不能代表全部地区和业务机构。第四，证据精确匹配只能证明输出可定位，"
        "不能替代生态学正确性判断。第五，0.6B模型强调低资源可复现性，结论不应外推至所有模型规模。",
    )

    add_paragraph(doc, "8 结论", style="Heading 1", first_indent=False)
    add_paragraph(
        doc,
        "本文重新界定了知识图谱在生态评估中的角色，提出面向技术规范数字化的EcoSpec-KG框架。"
        "该框架以HJ 1166-HJ 1176为语料，建立条款级来源标识、领域Schema、完整路径分组、双"
        "GraphRAG适配、LoRA补全及证据拒绝机制，并给出对比、消融、跨规范和专家验证方案。"
        "当前版本完成了数据与代码框架，但性能和专家结果仍待正式实验，因此本文暂不宣称方法优于"
        "基线。完成标注和GPU实验后，研究可为生态评估规范检索、参数依赖检查和标准知识复用提供"
        "可核验基础。",
    )
    add_paragraph(
        doc,
        "数据与代码可用性声明：项目代码、标注模板、数据卡和实验配置计划公开发布。正式标准PDF"
        "不随代码仓库重新分发，使用者需从权威渠道获取并通过环境变量配置本地语料目录。",
    )

    doc.add_page_break()
    add_paragraph(doc, "参考文献", style="Heading 1", first_indent=False)
    references = [
        "[1] 生态环境部. 全国生态状况调查评估技术规范—生态系统质量评估: HJ 1172—2021[S]. 北京: 生态环境部, 2021.",
        "[2] 生态环境部. 全国生态状况调查评估技术规范—生态系统服务功能评估: HJ 1173—2021[S]. 北京: 生态环境部, 2021.",
        "[3] 生态环境部. 全国生态状况调查评估技术规范—生态系统遥感解译与野外核查: HJ 1166—2021[S]. 北京: 生态环境部, 2021.",
        "[4] 生态环境部. 全国生态状况调查评估技术规范—数据质量控制与集成: HJ 1176—2021[S]. 北京: 生态环境部, 2021.",
        "[5] JI S, PAN S, CAMBRIA E, et al. A survey on knowledge graphs: representation, acquisition, and applications[J]. IEEE Transactions on Neural Networks and Learning Systems, 2022, 33(2): 494-514.",
        "[6] WANG Q, MAO Z, WANG B, et al. Knowledge graph embedding: a survey of approaches and applications[J]. IEEE Transactions on Knowledge and Data Engineering, 2017, 29(12): 2724-2743.",
        "[7] 郑晓云, 董仁才, 练岸鑫, 等. 基于多模态生态治理数据构建生态管理知识图谱技术[J]. 生态学报, 2024, 44(9): 3924-3933.",
        "[8] LEWIS P, PEREZ E, PIKTUS A, et al. Retrieval-augmented generation for knowledge-intensive NLP tasks[C]//Advances in Neural Information Processing Systems 33. 2020.",
        "[9] EDGE D, TRINH H, CHENG N, et al. From local to global: a graph RAG approach to query-focused summarization[EB/OL]. arXiv:2404.16130, 2024.",
        "[10] YANG A, LI A, YANG B, et al. Qwen3 technical report[EB/OL]. arXiv:2505.09388, 2025.",
        "[11] HU E J, SHEN Y, WALLIS P, et al. LoRA: low-rank adaptation of large language models[C]//International Conference on Learning Representations. 2022.",
        "[12] BORDES A, USUNIER N, GARCIA-DURÁN A, et al. Translating embeddings for modeling multi-relational data[C]//Advances in Neural Information Processing Systems 26. 2013.",
        "[13] YANG B, YIH W, HE X, et al. Embedding entities and relations for learning and inference in knowledge bases[C]//International Conference on Learning Representations. 2015.",
        "[14] TROUILLON T, WELBL J, RIEDEL S, et al. Complex embeddings for simple link prediction[C]//Proceedings of the 33rd International Conference on Machine Learning. 2016: 2071-2080.",
        "[15] YAO L, MAO C, LUO Y. KG-BERT: BERT for knowledge graph completion[EB/OL]. arXiv:1909.03193, 2019.",
        "[16] CHEN M, ZHANG W, ZHANG W, et al. Meta relational learning for few-shot link prediction in knowledge graphs[C]//EMNLP-IJCNLP. 2019: 4217-4226.",
        "[17] WANG L, ZHAO W, WEI Z, et al. SimKGC: simple contrastive knowledge graph completion with pre-trained language models[C]//Proceedings of ACL. 2022: 4281-4294.",
        "[18] CHEN C, WANG Y, LI B, et al. Knowledge is flat: a Seq2Seq generative framework for various knowledge graph completion[C]//Proceedings of COLING. 2022: 4005-4017.",
        "[19] YE H, ZHANG N, CHEN H, et al. Generative knowledge graph construction: a review[C]//Proceedings of EMNLP. 2022: 1-17.",
        "[20] TRAAG V A, WALTMAN L, VAN ECK N J. From Louvain to Leiden: guaranteeing well-connected communities[J]. Scientific Reports, 2019, 9: 5233.",
        "[21] BROWN T B, MANN B, RYDER N, et al. Language models are few-shot learners[C]//Advances in Neural Information Processing Systems 33. 2020: 1877-1901.",
        "[22] LEBO T, SAHOO S, MCGUINNESS D. PROV-O: the PROV ontology[EB/OL]. W3C Recommendation, 2013.",
        "[23] WILKINSON M D, DUMONTIER M, AALBERSBERG I J, et al. The FAIR guiding principles for scientific data management and stewardship[J]. Scientific Data, 2016, 3: 160018.",
        "[24] GEBRU T, MORGENSTERN J, VECCHIONE B, et al. Datasheets for datasets[J]. Communications of the ACM, 2021, 64(12): 86-92.",
        "[25] MICROSOFT. GraphRAG: a modular graph-based retrieval-augmented generation system[EB/OL]. Version 3.1.0, 2026.",
    ]
    for reference in references:
        style = "参考文献" if "参考文献" in doc.styles else None
        paragraph = add_paragraph(
            doc, reference, style=style, size=8, first_indent=False
        )
        paragraph.paragraph_format.keep_together = False
        paragraph.paragraph_format.line_spacing = 0.9
        paragraph.paragraph_format.space_after = Pt(0)

    # Ensure only one section and keep the template's page size.
    for extra in list(doc.sections)[1:]:
        extra.start_type = WD_SECTION.CONTINUOUS
    doc.core_properties.title = TITLE
    doc.core_properties.subject = "生态评估技术规范数字化与知识图谱"
    doc.core_properties.keywords = "生态评估;知识图谱;GraphRAG;证据溯源"
    doc.core_properties.author = "陈艳;沈志龙;张京;蒋娴"
    doc.save(DOCX_OUT)
    return DOCX_OUT


if __name__ == "__main__":
    print(build_document())
